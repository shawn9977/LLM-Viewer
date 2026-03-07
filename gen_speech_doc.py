from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 页面设置 ──────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(3)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── 样式辅助函数 ──────────────────────────────────────────
def set_run_font(run, size=11, bold=False, color=None, name_cn='微软雅黑', name_en='Calibri'):
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    run.font.name = name_en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name_cn)

def add_heading(doc, text, level=1, size=16, color=(31, 73, 125)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, color=color)
    return p

def add_subheading(doc, text, size=13, color=(68, 114, 196)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, color=color)
    return p

def add_body(doc, text, indent=False, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing = Pt(20)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    set_run_font(run, size=size)
    return p

def add_note(doc, text):
    """灰色备注/提示框"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run("【演讲提示】" + text)
    set_run_font(run, size=10, color=(128, 128, 128))
    return p

def add_highlight(doc, text, color=(192, 0, 0)):
    """重点强调段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.line_spacing = Pt(20)
    run = p.add_run(text)
    set_run_font(run, size=11, bold=True, color=color)
    return p

def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(1.2)
    p.paragraph_format.line_spacing = Pt(20)
    run = p.add_run(text)
    set_run_font(run, size=size)
    return p

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run("─" * 50)
    set_run_font(run, size=9, color=(180, 180, 180))
    return p

def add_table_simple(doc, headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=10, bold=True, color=(255,255,255))
        # blue background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F497D')
        tcPr.append(shd)
    # data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri+1]
        fill = 'EBF3FB' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                set_run_font(run, size=10)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tcPr.append(shd)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════
#  封面
# ══════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run("Qwen3-VL 多模态模型\nRoofline 性能分析")
set_run_font(run, size=22, bold=True, color=(31, 73, 125))

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(16)
run2 = p2.add_run("基于 LLM-Viewer 的三阶段推理性能建模")
set_run_font(run2, size=14, color=(68, 114, 196))

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(12)
run3 = p3.add_run("演讲稿  ·  约 15 分钟")
set_run_font(run3, size=11, color=(128, 128, 128))

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  演讲结构总览
# ══════════════════════════════════════════════════════════
add_heading(doc, "演讲结构总览", size=14, color=(31,73,125))
add_table_simple(doc,
    ["章节", "主题", "时长"],
    [
        ["第一部分", "为什么 VLM 需要单独分析？", "约 3 分钟"],
        ["第二部分", "VLM 架构拆解——三阶段计算流水线", "约 4 分钟"],
        ["第三部分", "Roofline 模型——三阶段性能瓶颈量化", "约 4 分钟"],
        ["第四部分", "LLM-Viewer 工具演示与优化洞察", "约 4 分钟"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  第一部分
# ══════════════════════════════════════════════════════════
add_heading(doc, "第一部分：为什么 VLM 需要单独分析？", size=15)
add_subheading(doc, "【开场白】")
add_body(doc,
    "大家好。今天我要和大家分享的主题是：如何用 Roofline 模型对多模态大语言模型——"
    "具体来说是 Qwen3-VL-8B 和 Qwen3-VL-32B——进行推理性能分析。"
)
add_body(doc,
    "在开始之前，我想先问大家一个问题：我们平时分析 LLM 的推理性能，"
    "用的是 prefill 和 decode 两个阶段。那么，当模型不只处理文字，还要处理图片的时候，"
    "这套框架还够用吗？"
)
add_highlight(doc, "答案是：不够用。这正是今天演讲的出发点。")

add_subheading(doc, "【三类模型的本质区别】")
add_body(doc,
    "我们先来看三类模型的本质区别。在 LLM-Viewer 中，我们支持三种模型类型："
)
add_bullet(doc, "LLM，比如 Qwen3-32B、Llama，输入是纯文本 Token，推理分为 prefill 和 decode 两个阶段，核心瓶颈是 KV Cache 的内存带宽。")
add_bullet(doc, "MoE，比如 Qwen3-MoE-30B，同样是纯文本输入，同样是两个阶段，但多了 Expert 路由和稀疏激活的开销。")
add_bullet(doc, "VLM，也就是今天的主角 Qwen3-VL，输入是图像加文本，推理分为三个阶段：vision、prefill、decode。核心瓶颈不只是 KV Cache，还有视觉编码器和跨模态对齐。")

add_body(doc,
    "这里最关键的一点是：VLM 是双分支架构。它由两个完全不同的神经网络组成——"
    "一个是 ViT 视觉编码器，负责把图片变成 Visual Token；"
    "另一个是 LLM 语言解码器，负责理解和生成文字。"
    "这两个分支的计算特征截然不同，不能用同一把尺子衡量，必须分开建模。"
)
add_note(doc, "此处可停顿，让听众消化三类模型的对比表格。")
add_divider(doc)

# ══════════════════════════════════════════════════════════
#  第二部分
# ══════════════════════════════════════════════════════════
add_heading(doc, "第二部分：VLM 架构拆解——三阶段计算流水线", size=15)
add_subheading(doc, "【三阶段流程讲解】")
add_body(doc,
    "接下来我们深入看 Qwen3-VL 的推理流程。整个推理过程分为三个串行阶段，"
    "我们逐一来看。"
)

add_subheading(doc, "Stage 1：Vision 阶段（VLM 专属）", size=12, color=(192,0,0))
add_body(doc,
    "第一阶段是 Vision 阶段，这是 VLM 相比 LLM 最核心的新增部分。"
    "输入一张 1024×1024 的图像，ViT 视觉编码器会把它切成一个个小 Patch。"
)
add_body(doc,
    "具体来说，patch_size 是 14 像素，所以横向和纵向各有 ceil(1024÷14) = 73 个 Patch，"
    "总共 73×73 = 5,329 个 Patch。每个 Patch 经过 ViT 的多层 Attention 和 FFN 处理后，"
    "再通过一个跨模态投影层 vision_proj，把视觉特征映射到语言模型的向量空间，"
    "变成 Visual Token。"
)
add_highlight(doc, "这一步在 LLM 和 MoE 中完全不存在，是 VLM 独有的计算开销。")

add_subheading(doc, "Stage 2：Prefill 阶段", size=12, color=(68,114,196))
add_body(doc,
    "第二阶段是 Prefill。这一阶段和纯 LLM 的 Prefill 逻辑相同，"
    "但有一个重要区别：序列长度更长。"
    "因为输入序列 = 文本 Token + Visual Token，"
    "一张 1024×1024 的图片就贡献了约 5,000 个 Token，"
    "这使得 Prefill 阶段的计算量远超同参数量的纯 LLM。"
)

add_subheading(doc, "Stage 3：Decode 阶段", size=12, color=(68,114,196))
add_body(doc,
    "第三阶段是 Decode，也就是自回归生成阶段。"
    "这一阶段和纯 LLM 完全相同——每一步生成一个 Token，"
    "反复执行直到生成结束符。"
)

add_subheading(doc, "【端到端指标：TTFT 与 TPOT】")
add_body(doc,
    "基于这三个阶段，我们定义两个端到端的性能指标："
)
add_bullet(doc, "TTFT（Time To First Token，首 Token 延迟）= vision 延迟 + prefill 延迟。这是用户发出请求到看到第一个字的等待时间。VLM 的 TTFT 远高于同参数 LLM，因为多了整个视觉编码阶段。")
add_bullet(doc, "TPOT（Time Per Output Token，每 Token 延迟）= decode 延迟。这是后续每生成一个字的时间，与纯 LLM 相同。")

add_subheading(doc, "【VLM 专属算子】")
add_body(doc,
    "Vision 阶段引入了一整套 LLM 完全没有的算子。"
    "在 LLM-Viewer 的分析中，这些算子会被单独列出并计算其延迟和内存访问量："
)
add_bullet(doc, "vision_patch_embed：将图像 Patch 转换为向量表示")
add_bullet(doc, "vision_qk_matmul / vision_sv_matmul：ViT 的 Attention 计算")
add_bullet(doc, "vision_gate/up/down_proj：ViT 的 FFN 层")
add_bullet(doc, "vision_norm1 / vision_norm2：ViT 的 LayerNorm")
add_bullet(doc, "vision_proj：跨模态投影，将视觉特征映射到语言空间")
add_note(doc, "此处可结合 LLM-Viewer 的算子图进行演示，指出视觉算子和文本算子的分布。")
add_divider(doc)

# ══════════════════════════════════════════════════════════
#  第三部分
# ══════════════════════════════════════════════════════════
add_heading(doc, "第三部分：Roofline 模型——三阶段性能瓶颈量化", size=15)
add_subheading(doc, "【Roofline 模型原理回顾】")
add_body(doc,
    "在进入 VLM 的具体分析之前，我先简单回顾一下 Roofline 模型的核心思想。"
)
add_body(doc,
    "Roofline 模型用一个指标来判断一个算子或阶段的性能瓶颈，这个指标叫做算术强度，"
    "定义为：计算量（OPs）除以内存访问量（Bytes）。"
    "同时，硬件有一个转折点，等于峰值算力除以内存带宽。"
)
add_body(doc,
    "判断规则很简单：如果算术强度低于转折点，说明内存带宽跟不上计算需求，"
    "这个阶段是 Memory-Bound，也就是带宽瓶颈；"
    "反之，如果算术强度高于转折点，说明算力是瓶颈，这个阶段是 Compute-Bound。"
)

add_subheading(doc, "【三阶段 Roofline 特征分析】")
add_body(doc,
    "现在我们用这个框架来分析 Qwen3-VL-32B 在 H100 上的三个阶段。"
)

add_subheading(doc, "Vision 阶段：Memory-Bound", size=12, color=(192,0,0))
add_body(doc,
    "Vision 阶段通常是 Memory-Bound。原因是 ViT 的层数相对较少，"
    "Patch 数量虽然有几千个，但整体计算密度不如长序列的 LLM Prefill。"
    "权重加载是主要开销，带宽利用率决定了这一阶段的性能上限。"
)

add_subheading(doc, "Prefill 阶段：Compute-Bound", size=12, color=(68,114,196))
add_body(doc,
    "Prefill 阶段在序列较长时是 Compute-Bound。"
    "由于 Visual Token 的加入，序列长度可能达到 5,000 到 10,000，"
    "大量的矩阵乘法使得算术强度很高，算力成为瓶颈。"
    "这一点和纯 LLM 的 Prefill 类似，但序列更长，计算量更大。"
)

add_subheading(doc, "Decode 阶段：Memory-Bound", size=12, color=(68,114,196))
add_body(doc,
    "Decode 阶段是典型的 Memory-Bound，这和纯 LLM 完全相同。"
    "每一步只生成一个 Token，权重的复用率极低，"
    "每次都要把几十 GB 的模型权重从显存加载一遍，带宽是绝对瓶颈。"
)

add_subheading(doc, "【与纯 LLM 的对比——重点】")
add_highlight(doc,
    "这里是今天演讲最重要的对比点，请大家注意：",
    color=(31,73,125)
)
add_body(doc,
    "纯 LLM 只有两个阶段：Compute-Bound 的 Prefill 和 Memory-Bound 的 Decode。"
    "而 VLM 多了一个 Memory-Bound 的 Vision 阶段。"
    "这意味着 VLM 有两个 Memory-Bound 阶段，带宽压力更大，"
    "而且 Vision 阶段的带宽消耗和 Decode 阶段的带宽消耗是完全不同的来源——"
    "前者来自 ViT 权重，后者来自 LLM 权重和 KV Cache。"
)

add_subheading(doc, "【图像分辨率的影响——VLM 独有变量】")
add_body(doc,
    "VLM 有一个纯 LLM 完全没有的性能调优变量：图像分辨率。"
)
add_body(doc,
    "以 patch_size=14 为例：1024×1024 的图像产生约 5,329 个 Visual Token，"
    "而 512×512 的图像只产生约 1,369 个 Visual Token，减少了 75%。"
    "这不仅直接降低了 Vision 阶段的计算量，"
    "还缩短了 Prefill 阶段的序列长度，对 TTFT 有双重优化效果。"
)
add_highlight(doc, "图像分辨率是 VLM 性能调优的第一旋钮，优先于量化和并行策略。")
add_note(doc, "此处可展示 LLM-Viewer 中切换 image_size 参数后 Roofline 图的变化。")
add_divider(doc)

# ══════════════════════════════════════════════════════════
#  第四部分
# ══════════════════════════════════════════════════════════
add_heading(doc, "第四部分：LLM-Viewer 工具演示与优化洞察", size=15)
add_subheading(doc, "【工具能力介绍】")
add_body(doc,
    "LLM-Viewer 对 Qwen3-VL 的支持是完整的三阶段独立建模。"
    "在工具中选择 Qwen3-VL-8B 或 32B，配置硬件和分析参数后，"
    "工具会分别输出 vision、prefill、decode 三个阶段的 Roofline 图，"
    "以及逐算子的延迟分解、TTFT/TPOT 端到端估算和内存占用分析。"
)
add_body(doc,
    "特别值得注意的是 image_size 参数——只有设置了这个参数，"
    "工具才会触发视觉分支的分析。如果不设置，工具只分析文本部分，"
    "这对于理解 VLM 的完整性能画像是不够的。"
)

add_subheading(doc, "【关键优化手段】")
add_body(doc, "基于 Roofline 分析，我们可以得出以下优化建议，按优先级排列：")
add_bullet(doc, "第一优先：降低图像分辨率或增大 spatial_merge_size。这是成本最低、效果最显著的优化，直接减少 Visual Token 数量，同时优化 Vision 和 Prefill 两个阶段。")
add_bullet(doc, "第二优先：权重量化（w_bit=4）。对于 Memory-Bound 的 Vision 和 Decode 阶段，减少权重加载量可以直接提升带宽利用率，效果非常显著。")
add_bullet(doc, "第三优先：张量并行（TP=4 或 8）。LLM-Viewer 对视觉分支和文本分支都支持 TP，可以线性扩展算力，适合对延迟要求极高的场景。")
add_bullet(doc, "第四优先：Flash Attention。视觉分支和文本分支都可以开启，减少 Attention 的内存访问量，对长序列效果更明显。")

add_subheading(doc, "【8B vs 32B 的选型建议】")
add_body(doc,
    "最后，关于 8B 和 32B 的选型，我们可以从 Roofline 分析中得到量化依据。"
)
add_body(doc,
    "Qwen3-VL-8B 和 32B 的视觉编码器参数量相近——ViT 的规模相对独立于语言模型的规模。"
    "因此两者的 Vision 阶段延迟差异不大。"
    "主要差异在 Decode 阶段：32B 的权重更大，Decode 阶段的内存带宽压力更高，"
    "TPOT 更长。如果对话场景对实时性要求高，8B 是更好的选择；"
    "如果需要更强的多模态理解能力，32B 配合量化或 TP 是推荐方案。"
)
add_note(doc, "此处可在 LLM-Viewer 中切换 8B 和 32B，对比两者的 Decode 阶段 Roofline 图。")
add_divider(doc)

# ══════════════════════════════════════════════════════════
#  总结
# ══════════════════════════════════════════════════════════
add_heading(doc, "总结", size=15)
add_body(doc,
    "今天的演讲，我们从三个层次分析了 Qwen3-VL 的推理性能："
)
add_bullet(doc, "架构层：VLM 是双分支架构，Vision 阶段是 LLM 完全没有的新增计算，必须单独建模。")
add_bullet(doc, "性能层：三阶段各有不同的 Roofline 特征——Vision 是 Memory-Bound，Prefill 是 Compute-Bound，Decode 是 Memory-Bound。")
add_bullet(doc, "优化层：图像分辨率是 VLM 独有的第一调优旋钮，其次是量化、TP 和 Flash Attention。")

add_highlight(doc,
    "核心结论：VLM 的性能分析不能简单套用 LLM 的框架。"
    "LLM-Viewer 通过三阶段独立 Roofline 建模，"
    "为多模态模型的硬件选型和推理优化提供了量化依据。",
    color=(31,73,125)
)
add_body(doc, "谢谢大家，欢迎提问。")

# ── 保存 ──────────────────────────────────────────────────
doc.save('/data/project/LLM-Viewer/演讲稿_Qwen3VL_Roofline分析.docx')
print("Done")
